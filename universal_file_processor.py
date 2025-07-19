import os
import json
import re
import tiktoken
import chardet
from pathlib import Path
from typing import List, Dict, Any

# Simple text splitter implementation
class SimpleTextSplitter:
    def __init__(self, chunk_size=800, chunk_overlap=100, separators=None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", ". ", " ", ""]
    
    def split_text(self, text):
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        current_chunk = ""
        
        # Split by separators in order of preference
        for separator in self.separators:
            if separator in text:
                parts = text.split(separator)
                for part in parts:
                    if len(current_chunk) + len(part) + len(separator) <= self.chunk_size:
                        if current_chunk:
                            current_chunk += separator + part
                        else:
                            current_chunk = part
                    else:
                        if current_chunk:
                            chunks.append(current_chunk)
                        current_chunk = part
                        
                        # If single part is too long, split it
                        while len(current_chunk) > self.chunk_size:
                            chunks.append(current_chunk[:self.chunk_size])
                            current_chunk = current_chunk[self.chunk_size - self.chunk_overlap:]
                
                if current_chunk:
                    chunks.append(current_chunk)
                return chunks
        
        # Fallback: simple character-based splitting
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - self.chunk_overlap
        
        return chunks

from openai import OpenAI
from supabase import create_client

class UniversalFileProcessor:
    def __init__(self):
        self.supported_extensions = {'.txt', '.pdf', '.doc', '.docx', '.csv'}
        self.chunks = []
        # Initialize with None - will be set by web server with user credentials
        self.openai_api_key = None
        self.supabase_url = None
        self.supabase_service_key = None
    
    def get_openai_client(self):
        """Get OpenAI client with user-provided API key"""
        if not self.openai_api_key:
            raise ValueError("OpenAI API key not provided")
        return OpenAI(api_key=self.openai_api_key)
    
    def get_supabase_client(self):
        """Get Supabase client with user-provided credentials"""
        if not self.supabase_url or not self.supabase_service_key:
            raise ValueError("Supabase credentials not provided")
        
        # Create client with minimal configuration to avoid proxy issues
        try:
            return create_client(
                supabase_url=self.supabase_url,
                supabase_key=self.supabase_service_key
            )
        except Exception as e:
            # Fallback for older versions
            return create_client(self.supabase_url, self.supabase_service_key)
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        return text.strip()
        
    def extract_text_from_file(self, file_path: str, file_content: bytes = None) -> str:
        """Extract text from various file formats"""
        file_extension = Path(file_path).suffix.lower()
        
        print(f"Processing file: {file_path} with extension: {file_extension}")
        
        try:
            if file_extension == '.txt':
                if file_content:
                    # Detect encoding
                    detected = chardet.detect(file_content)
                    encoding = detected.get('encoding', 'utf-8') if detected else 'utf-8'
                    print(f"Detected encoding: {encoding}")
                    try:
                        text = file_content.decode(encoding)
                    except (UnicodeDecodeError, LookupError):
                        print("Fallback to UTF-8 with error handling")
                        text = file_content.decode('utf-8', errors='ignore')
                else:
                    # Detect encoding for file
                    with open(file_path, 'rb') as f:
                        raw_data = f.read()
                        detected = chardet.detect(raw_data)
                        encoding = detected.get('encoding', 'utf-8') if detected else 'utf-8'
                    
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            text = file.read()
                    except (UnicodeDecodeError, LookupError):
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                            text = file.read()
                
                print(f"Extracted {len(text)} characters from TXT file")
                return self.clean_text(text)
            
            elif file_extension == '.pdf':
                import PyPDF2
                
                try:
                    if file_content:
                        import io
                        pdf_file = io.BytesIO(file_content)
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                    else:
                        with open(file_path, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                    
                    text_parts = []
                    for i, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_parts.append(f"Page {i+1}:\n{page_text}")
                        except Exception as e:
                            print(f"Error extracting page {i+1}: {e}")
                            continue
                    
                    text = "\n\n".join(text_parts)
                    print(f"Extracted {len(text)} characters from PDF file")
                    return self.clean_text(text)
                except Exception as e:
                    raise Exception(f"Error processing PDF file: {e}")
            
            elif file_extension in ['.doc', '.docx']:
                from docx import Document
                
                try:
                    if file_content:
                        import io
                        doc_file = io.BytesIO(file_content)
                        doc = Document(doc_file)
                    else:
                        doc = Document(file_path)
                    
                    text_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                    
                    text = "\n\n".join(text_parts)
                    print(f"Extracted {len(text)} characters from Word document")
                    return self.clean_text(text)
                except Exception as e:
                    raise Exception(f"Error processing Word document: {e}")
            
            elif file_extension == '.csv':
                import pandas as pd
                
                try:
                    if file_content:
                        import io
                        csv_file = io.StringIO(file_content.decode('utf-8', errors='ignore'))
                        df = pd.read_csv(csv_file)
                    else:
                        df = pd.read_csv(file_path)
                    
                    # Convert DataFrame to structured text
                    text_parts = []
                    text_parts.append(f"CSV Data with {len(df)} rows and {len(df.columns)} columns")
                    text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
                    text_parts.append("")
                    
                    # Add sample data
                    for idx, row in df.head(100).iterrows():  # Limit to first 100 rows
                        row_text = f"Row {idx + 1}: "
                        row_items = []
                        for col, val in row.items():
                            if pd.notna(val) and str(val).strip():
                                row_items.append(f"{col}: {val}")
                        if row_items:
                            row_text += "; ".join(row_items)
                            text_parts.append(row_text)
                    
                    text = "\n".join(text_parts)
                    print(f"Extracted {len(text)} characters from CSV file")
                    return self.clean_text(text)
                except Exception as e:
                    raise Exception(f"Error processing CSV file: {e}")
            
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            print(f"Error in extract_text_from_file: {e}")
            raise Exception(f"Error extracting text from {Path(file_path).name}: {e}")
    
    def count_tokens(self, text: str, model: str = "gpt-3.5-turbo") -> int:
        """Count tokens in text using tiktoken"""
        try:
            encoding = tiktoken.encoding_for_model(model)
            return len(encoding.encode(text))
        except Exception as e:
            print(f"Error counting tokens: {e}")
            # Fallback: rough estimation (1 token ≈ 4 characters)
            return len(text) // 4
    
    def create_chunks(self, text: str, source: str, title: str = None) -> List[Dict[str, Any]]:
        """Create chunks from text with metadata"""
        if not text.strip():
            print("Warning: Empty text provided for chunking")
            return []
        
        print(f"Creating chunks from {len(text)} characters of text")
        
        # Initialize text splitter
        text_splitter = SimpleTextSplitter(
            chunk_size=800,  # Smaller chunks for better embedding quality
            chunk_overlap=100,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Split text into chunks
        chunks = text_splitter.split_text(text)
        print(f"Text splitter created {len(chunks)} raw chunks")
        
        # Create chunk objects with metadata
        chunk_objects = []
        for i, chunk in enumerate(chunks):
            if chunk.strip():  # Only add non-empty chunks
                # Clean the chunk
                clean_chunk = self.clean_text(chunk)
                if len(clean_chunk) < 20:  # Skip very short chunks
                    continue
                
                chunk_obj = {
                    "content": clean_chunk,
                    "source": source,
                    "title": title or Path(source).stem,
                    "chunk_index": i,
                    "chunk_type": "text",
                    "token_count": self.count_tokens(clean_chunk),
                    "metadata": {
                        "file_extension": Path(source).suffix.lower(),
                        "total_chunks": len(chunks),
                        "char_count": len(clean_chunk)
                    }
                }
                chunk_objects.append(chunk_obj)
        
        print(f"Created {len(chunk_objects)} valid chunk objects")
        return chunk_objects
    
    def get_embedding(self, text: str) -> List[float]:
        """Generate embedding for text"""
        try:
            client = self.get_openai_client()
            
            # Truncate text if too long for embedding model
            max_length = 8000
            if len(text) > max_length:
                text = text[:max_length]
            
            response = client.embeddings.create(
                input=text,
                model="text-embedding-3-small"
            )
            return response.data[0].embedding
        except Exception as e:
            raise Exception(f"Error generating embedding: {e}")
    
    def process_file(self, file_path: str, file_content: bytes = None) -> List[Dict[str, Any]]:
        """Process a single file and return chunks"""
        print(f"Processing file: {file_path}")
        
        try:
            # Extract text from file
            text = self.extract_text_from_file(file_path, file_content)
            
            if not text.strip():
                raise ValueError(f"No text could be extracted from {file_path}")
            
            # Create chunks
            chunks = self.create_chunks(text, file_path)
            
            if not chunks:
                raise ValueError(f"No valid chunks could be created from {file_path}")
            
            print(f"Created {len(chunks)} chunks from {file_path}")
            return chunks
            
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            raise
    
    def upload_to_supabase(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Upload all chunks to Supabase with embeddings"""
        print("Starting upload to Supabase...")
        
        try:
            # Get user's Supabase client
            supabase = self.get_supabase_client()
            
            # Test connection
            health_check = supabase.table("documents").select("id").limit(1).execute()
            print("Supabase connection verified")
            
        except Exception as e:
            raise Exception(f"Failed to connect to Supabase: {e}")
        
        successful_uploads = 0
        failed_uploads = 0
        embedding_errors = 0
        
        for i, chunk in enumerate(chunks, 1):
            try:
                # Generate embedding
                try:
                    embedding = self.get_embedding(chunk["content"])
                except Exception as e:
                    print(f"Chunk {i}: Failed to generate embedding - {e}")
                    embedding_errors += 1
                    continue
                
                # Prepare data for Supabase
                data = {
                    "content": chunk["content"],
                    "embedding": embedding,
                    "source": chunk["source"],
                    "metadata": {
                        "title": chunk["title"],
                        "chunk_type": chunk["chunk_type"],
                        **chunk.get("metadata", {})
                    }
                }
                
                # Insert into Supabase
                response = supabase.table("documents").insert(data).execute()
                
                if hasattr(response, 'error') and response.error:
                    print(f"Chunk {i}: Upload failed - {response.error}")
                    failed_uploads += 1
                else:
                    successful_uploads += 1
                    if i % 10 == 0:  # Progress update every 10 chunks
                        print(f"Uploaded {i}/{len(chunks)} chunks...")
                
            except Exception as e:
                print(f"Chunk {i}: Exception during upload - {e}")
                failed_uploads += 1
        
        result = {
            "successful_uploads": successful_uploads,
            "failed_uploads": failed_uploads,
            "embedding_errors": embedding_errors,
            "total_chunks": len(chunks)
        }
        
        print(f"Upload complete! Success: {successful_uploads}, Failed: {failed_uploads}, Embedding errors: {embedding_errors}")
        return result

def main():
    """Main function for testing"""
    print("This is a library module. Use the web interface to process files.")
    print("Start the web server with: python web_server.py")

if __name__ == "__main__":
    main()